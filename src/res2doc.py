# %%
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import json
import os

def set_paragraph_font(paragraph, font_size=12, chinese_font='宋体', english_font='Times New Roman'):
    """设置段落字体格式：小四号，汉字宋体，英文Times New Roman"""
    # 如果段落没有runs，先添加一个空的run
    if not paragraph.runs:
        paragraph.add_run('')

    for run in paragraph.runs:
        run.font.size = Pt(font_size)
        run.font.name = english_font

        # 使用更完整的字体设置方法
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), english_font)
        rFonts.set(qn('w:hAnsi'), english_font)
        rFonts.set(qn('w:eastAsia'), chinese_font)
        rFonts.set(qn('w:cs'), chinese_font)

def set_heading_font(heading, font_size=16, font_name='黑体', color='000000'):
    """设置标题字体格式和颜色"""
    if not heading.runs:
        heading.add_run('')

    for run in heading.runs:
        run.font.size = Pt(font_size)
        run.font.name = font_name

        # 使用更完整的字体设置方法
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:cs'), font_name)

        # 设置颜色
        color_elem = rPr.get_or_add_color()
        color_elem.set(qn('w:val'), color)

def res_to_doc(sjt_data, output_docx_file=None):
    """生成SJT文档的主函数

    Args:
        sjt_data (dict): SJT数据字典
        output_docx_file (str, optional): 输出的DOCX文件路径，如果不指定则自动生成

    Returns:
        str: 生成的文档路径，如果失败返回None
    """
    # 检查数据格式
    if not isinstance(sjt_data, dict):
        print("错误：输入数据必须是字典格式")
        return None

    # 生成文件名
    if output_docx_file is None:
        timestamp = datetime.now().strftime('%Y%m%d')
        base_name = f'SJTAgent_v0.1_{timestamp}'
        extension = '.docx'
        docx_path = f'{base_name}{extension}'

        # 避免文件名冲突
        counter = 1
        while os.path.exists(docx_path):
            docx_path = f'{base_name}_{counter}{extension}'
            counter += 1
    else:
        docx_path = output_docx_file

    doc = Document()

    # 设置大标题：居中对齐，黑体三号
    title = doc.add_heading('SJTAgent-text生成结果', level=1)
    set_heading_font(title, font_size=16, font_name='黑体')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for trait_key in sjt_data:
        # 设置小标题：黑体小四号
        subtitle = doc.add_heading(f'特质：{trait_key}', level=2)
        set_heading_font(subtitle, font_size=12, font_name='黑体')

        trait_items = sjt_data[trait_key]
        for idx_str in sorted(trait_items.keys(), key=lambda x: int(x)):
            sjt = trait_items[idx_str]

            # 题目段落：小四号，汉字宋体，英文Times New Roman
            title_para = doc.add_paragraph(f'题目 {idx_str}（特质：{trait_key}）')
            set_paragraph_font(title_para)

            if isinstance(sjt, dict) and 'situation' in sjt:
                # 情景段落：小四号，汉字宋体，英文Times New Roman
                situation_para = doc.add_paragraph(f'情景：{sjt["situation"]}')
                set_paragraph_font(situation_para)

            options = sjt.get('options') if isinstance(sjt, dict) else None
            if isinstance(options, dict):
                for opt_key in ['A', 'B', 'C', 'D']:
                    if opt_key in options:
                        # 选项段落：小四号，汉字宋体，英文Times New Roman
                        option_para = doc.add_paragraph(f'{opt_key}. {options[opt_key]}')
                        set_paragraph_font(option_para)

            # 空行
            empty_para = doc.add_paragraph('')
            if empty_para.runs:
                set_paragraph_font(empty_para)

    doc.save(docx_path)
    print(f"文档已保存为: {docx_path}")
    return docx_path

# 直接执行时运行
if __name__ == "__main__":
    with open('output/sjt-text.json', 'r', encoding='utf-8') as f:
        data = json.load(f)
    res_to_doc(data)
